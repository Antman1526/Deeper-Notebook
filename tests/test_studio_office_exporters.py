"""Semantic and trust-boundary contracts for editable Studio Office exports."""

from __future__ import annotations

import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from open_notebook.domain.notebook import StudioArtifact
from open_notebook.studio.exporters import export_document, export_spreadsheet
from open_notebook.studio.generation.persistence import persist_artifact_exports
from open_notebook.studio.payloads import build_structured_payload
from open_notebook.studio.schemas import (
    CoursePackDocument,
    DataTableDocument,
    GenericDocument,
    ResearchRunDocument,
    parse_artifact_document,
)


def _archive_members(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as package:
        return package.namelist()


def _assert_safe_office_archive(path: Path) -> None:
    with zipfile.ZipFile(path) as package:
        members = package.namelist()
        assert not any(
            "vbaProject" in member or member.endswith(".bin") for member in members
        )
        for member in members:
            if not member.endswith(".rels"):
                continue
            assert b'TargetMode="External"' not in package.read(member)


def _report() -> GenericDocument:
    document = parse_artifact_document(
        "report",
        {
            "artifact_type": "report",
            "title": "Local Evidence Report",
            "summary": "A source-grounded report for a private research notebook.",
            "sections": [
                {
                    "heading": "Findings",
                    "body": "The reviewed sources agree on the core workflow.",
                    "bullets": [
                        "Preserve source ownership.",
                        "Show citations in context.",
                    ],
                    "citations": ["[S1]", "[S2]"],
                }
            ],
        },
    )
    assert isinstance(document, GenericDocument)
    return document


def _course_pack() -> CoursePackDocument:
    document = parse_artifact_document(
        "course_pack",
        {
            "artifact_type": "course_pack",
            "title": "Evidence Studio Course Pack",
            "audience": "Private researchers",
            "learning_outcomes": ["Create an editable, source-grounded export."],
            "prerequisites": ["A selected source set."],
            "modules": [
                {
                    "title": "Module 1: Trusted outputs",
                    "summary": "Build documents that remain editable and auditable.",
                    "lessons": [
                        {
                            "title": "Export review",
                            "content": "Open the file and validate citations before sharing.",
                            "duration_minutes": 20,
                            "exercise": "Inspect the citation appendix.",
                            "facilitator_notes": "Demonstrate the local-only workflow.",
                            "citations": ["[S1]"],
                        }
                    ],
                }
            ],
            "final_assessment": [
                {
                    "prompt": "Which output remains editable?",
                    "options": [
                        {"id": "a", "text": "DOCX"},
                        {"id": "b", "text": "A screenshot"},
                    ],
                    "correct_option_id": "a",
                    "explanation": "DOCX is an editable Office document.",
                    "citations": ["[S1]"],
                }
            ],
        },
    )
    assert isinstance(document, CoursePackDocument)
    return document


def _research_run() -> ResearchRunDocument:
    document = parse_artifact_document(
        "research_run",
        {
            "artifact_type": "research_run",
            "title": "Model routing research",
            "objective": "Compare local model routing receipts.",
            "hypotheses": ["Quality-aware selection improves useful outputs."],
            "stages": [
                {
                    "title": "Evidence review",
                    "findings": [
                        {
                            "text": "The benchmark keeps route decisions inspectable.",
                            "citations": ["[S3]"],
                        }
                    ],
                    "status": "complete",
                }
            ],
            "gaps": ["A native package smoke still needs manual review."],
            "next_actions": ["Open the export in an Office editor."],
        },
    )
    assert isinstance(document, ResearchRunDocument)
    return document


def _table() -> DataTableDocument:
    document = parse_artifact_document(
        "data_table",
        {
            "artifact_type": "data_table",
            "title": "Evaluation measurements",
            "columns": ["Date", "Score", "Note"],
            "rows": [
                {
                    "values": {
                        "Date": "2026-07-17",
                        "Score": "95.5",
                        "Note": "=not-a-formula",
                    },
                    "citations": ["[S1]"],
                },
                {
                    "values": {
                        "Date": "2026-07-18",
                        "Score": "88",
                        "Note": "Reviewed locally",
                    },
                    "citations": ["[S2]"],
                },
            ],
        },
    )
    assert isinstance(document, DataTableDocument)
    return document


def test_docx_exports_structured_documents_with_properties_tables_and_citations(
    tmp_path: Path,
) -> None:
    path = tmp_path / "report.docx"

    export_document(_report(), path)

    reopened = Document(path)
    text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    assert "Local Evidence Report" in text
    assert "Findings" in text
    assert "[S1] [S2]" in text
    assert reopened.core_properties.title == "Local Evidence Report"
    assert reopened.core_properties.author == "Open Notebook Plus"
    assert reopened.tables, "citation appendix should be a real Word table"
    assert _archive_members(path)[0] == "[Content_Types].xml"
    _assert_safe_office_archive(path)


def test_docx_supports_course_packs_and_research_runs_with_page_boundaries(
    tmp_path: Path,
) -> None:
    course_path = tmp_path / "course-pack.docx"
    research_path = tmp_path / "research-run.docx"

    export_document(_course_pack(), course_path)
    export_document(_research_run(), research_path)

    course = Document(course_path)
    course_text = "\n".join(paragraph.text for paragraph in course.paragraphs)
    assert "Learning outcomes" in course_text
    assert "Module 1: Trusted outputs" in course_text
    assert "Assessment" in course_text
    assert len(course.tables) >= 2
    assert any('w:type="page"' in paragraph._p.xml for paragraph in course.paragraphs)

    research = Document(research_path)
    research_text = "\n".join(paragraph.text for paragraph in research.paragraphs)
    assert "Evidence review" in research_text
    assert "[S3]" in research_text
    _assert_safe_office_archive(course_path)
    _assert_safe_office_archive(research_path)


def test_xlsx_exports_typed_rows_source_markers_and_validated_chart(
    tmp_path: Path,
) -> None:
    path = tmp_path / "measurements.xlsx"

    export_spreadsheet(_table(), path)

    workbook = load_workbook(path, data_only=False)
    sheet = workbook.active
    assert sheet.title == "Evaluation measurements"
    assert sheet.freeze_panes == "A2"
    assert sheet.auto_filter.ref == "A1:D3"
    assert [cell.value for cell in sheet[1]] == [
        "Date",
        "Score",
        "Note",
        "Source markers",
    ]
    assert isinstance(sheet["A2"].value, date)
    assert sheet["A2"].value.isoformat().startswith("2026-07-17")
    assert sheet["B2"].value == 95.5
    assert sheet["C2"].value == "'=not-a-formula"
    assert sheet["C2"].data_type == "s"
    assert sheet["D2"].value == "[S1]"
    assert sheet._charts
    assert all(dimension.width <= 60 for dimension in sheet.column_dimensions.values())
    assert not any(cell.data_type == "f" for row in sheet.iter_rows() for cell in row)
    _assert_safe_office_archive(path)


def test_persistence_attaches_office_exports_for_validated_documents(
    tmp_path: Path, monkeypatch
) -> None:
    monkeypatch.setenv("OPEN_NOTEBOOK_ARTIFACT_EXPORT_DIR", str(tmp_path))
    report = _report()
    artifact = StudioArtifact(
        id="studio_artifact:office-report",
        notebook_id="notebook:private",
        artifact_type="report",
        title=report.title,
        status="completed",
        output_payload=build_structured_payload(report, "# Local Evidence Report"),
    )

    paths = persist_artifact_exports(artifact, "# Local Evidence Report")

    assert Path(paths["docx"]).exists()
    assert artifact.export_paths["docx"] == paths["docx"]
    assert Path(paths["json"]).exists()
