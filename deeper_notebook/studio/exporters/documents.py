"""Trusted DOCX exports for structured Evidence Studio documents.

The document model is parsed before this module is called.  This exporter only
maps those typed values to Word primitives; it never accepts model-authored
Office XML, relationships, macros, or links.
"""

from __future__ import annotations

from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

from deeper_notebook.identity import PRODUCT_NAME
from deeper_notebook.studio.exporters.metadata import brand_office_application
from deeper_notebook.studio.schemas import (
    ArtifactDocumentBase,
    CoursePackDocument,
    GenericDocument,
    ResearchRunDocument,
)

_EDITABLE_DOCUMENTS = (GenericDocument, CoursePackDocument, ResearchRunDocument)


def _add_paragraph(document: Document, text: str, *, style: str | None = None) -> None:
    if text.strip():
        document.add_paragraph(text.strip(), style=style)


def _add_citations(document: Document, citations: Iterable[str]) -> None:
    markers = " ".join(dict.fromkeys(citations))
    if markers:
        _add_paragraph(document, f"Sources  {markers}", style="Evidence Citation")


def _configure_document(document: Document, title: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    citation = document.styles.add_style("Evidence Citation", WD_STYLE_TYPE.PARAGRAPH)
    citation.font.name = "Arial"
    citation.font.size = Pt(9)
    citation.font.italic = True

    properties = document.core_properties
    properties.title = title
    properties.author = PRODUCT_NAME
    properties.last_modified_by = PRODUCT_NAME
    properties.subject = "Evidence Studio editable export"
    properties.keywords = "evidence, local research, editable export"
    properties.comments = (
        "Generated locally from a structured, source-grounded artifact."
    )


def _add_citation_appendix(document: Document, markers: Iterable[str]) -> None:
    unique_markers = list(dict.fromkeys(markers))
    if not unique_markers:
        return

    document.add_page_break()
    document.add_heading("Citation appendix", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Marker"
    table.rows[0].cells[1].text = "Export note"
    for marker in unique_markers:
        cells = table.add_row().cells
        cells[0].text = marker
        cells[
            1
        ].text = "Stored source marker; review in the notebook for source context."


def _export_generic(document: Document, artifact: GenericDocument) -> list[str]:
    citations: list[str] = []
    _add_paragraph(document, artifact.summary)
    for section in artifact.sections:
        document.add_heading(section.heading, level=1)
        _add_paragraph(document, section.body)
        for bullet in section.bullets:
            _add_paragraph(document, bullet, style="List Bullet")
        _add_citations(document, section.citations)
        citations.extend(section.citations)
    return citations


def _export_course_pack(document: Document, artifact: CoursePackDocument) -> list[str]:
    citations: list[str] = []
    _add_paragraph(document, f"Audience: {artifact.audience}")
    document.add_heading("Learning outcomes", level=1)
    outcome_table = document.add_table(rows=1, cols=1)
    outcome_table.style = "Light Shading Accent 1"
    outcome_table.rows[0].cells[0].text = "Learners will be able to"
    for outcome in artifact.learning_outcomes:
        outcome_table.add_row().cells[0].text = outcome
    if artifact.prerequisites:
        document.add_heading("Prerequisites", level=1)
        for prerequisite in artifact.prerequisites:
            _add_paragraph(document, prerequisite, style="List Bullet")

    for index, module in enumerate(artifact.modules, start=1):
        document.add_page_break()
        document.add_heading(module.title, level=1)
        _add_paragraph(document, module.summary)
        for lesson in module.lessons:
            document.add_heading(lesson.title, level=2)
            if lesson.duration_minutes:
                _add_paragraph(
                    document, f"Estimated duration: {lesson.duration_minutes} minutes"
                )
            _add_paragraph(document, lesson.content)
            if lesson.exercise:
                _add_paragraph(
                    document, f"Exercise: {lesson.exercise}", style="List Bullet"
                )
            if lesson.facilitator_notes:
                _add_paragraph(
                    document, f"Facilitator notes: {lesson.facilitator_notes}"
                )
            _add_citations(document, lesson.citations)
            citations.extend(lesson.citations)

    if artifact.final_assessment:
        document.add_page_break()
        document.add_heading("Assessment", level=1)
        assessment = document.add_table(rows=1, cols=3)
        assessment.style = "Light Shading Accent 1"
        header = assessment.rows[0].cells
        header[0].text = "Question"
        header[1].text = "Options"
        header[2].text = "Explanation"
        for question in artifact.final_assessment:
            cells = assessment.add_row().cells
            cells[0].text = question.prompt
            cells[1].text = "\n".join(option.text for option in question.options)
            cells[2].text = question.explanation
            _add_citations(document, question.citations)
            citations.extend(question.citations)
    return citations


def _export_research_run(
    document: Document, artifact: ResearchRunDocument
) -> list[str]:
    citations: list[str] = []
    document.add_heading("Objective", level=1)
    _add_paragraph(document, artifact.objective)
    if artifact.hypotheses:
        document.add_heading("Hypotheses", level=1)
        for hypothesis in artifact.hypotheses:
            _add_paragraph(document, hypothesis, style="List Bullet")
    document.add_page_break()
    document.add_heading("Research stages", level=1)
    for stage in artifact.stages:
        document.add_heading(stage.title, level=2)
        _add_paragraph(document, f"Status: {stage.status}")
        findings = document.add_table(rows=1, cols=2)
        findings.style = "Light Shading Accent 1"
        findings.rows[0].cells[0].text = "Finding"
        findings.rows[0].cells[1].text = "Sources"
        for finding in stage.findings:
            cells = findings.add_row().cells
            cells[0].text = finding.text
            cells[1].text = " ".join(finding.citations)
            _add_citations(document, finding.citations)
            citations.extend(finding.citations)
    if artifact.gaps:
        document.add_heading("Gaps", level=1)
        for gap in artifact.gaps:
            _add_paragraph(document, gap, style="List Bullet")
    if artifact.next_actions:
        document.add_heading("Next actions", level=1)
        for action in artifact.next_actions:
            _add_paragraph(document, action, style="List Bullet")
    return citations


def export_document(artifact: ArtifactDocumentBase, path: Path) -> None:
    """Write a safe, editable DOCX for a structured textual Studio artifact."""
    if not isinstance(artifact, _EDITABLE_DOCUMENTS):
        raise TypeError(
            "export_document requires GenericDocument, CoursePackDocument, or "
            "ResearchRunDocument"
        )
    if path.suffix.lower() != ".docx":
        raise ValueError("DOCX export path must end in .docx")

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document, artifact.title)
    document.add_heading(artifact.title, level=0)
    _add_paragraph(document, f"{PRODUCT_NAME} / Evidence Studio")

    if isinstance(artifact, GenericDocument):
        citations = _export_generic(document, artifact)
    elif isinstance(artifact, CoursePackDocument):
        citations = _export_course_pack(document, artifact)
    else:
        citations = _export_research_run(document, artifact)

    _add_citation_appendix(document, citations)
    document.save(path)
    brand_office_application(path)


__all__ = ["export_document"]
